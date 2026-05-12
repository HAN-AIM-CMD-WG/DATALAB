"""DOCX-extractie en -herbouw via python-docx.

We lopen over alle paragrafen (body, table-cellen, headers, footers) en
behandelen iedere paragraaf als één blok. Vervangingen worden per
paragraaf in één keer toegepast op de gecombineerde tekst en daarna
teruggeschreven naar één enkele plain run (alle eerdere runs en
``<w:hyperlink>``-children worden gewist). Dat is destructief voor
inline-opmaak (bold midden in een zin, mailto-links) maar voorkomt twee
klassen van bugs die we eerder hadden:

1. **Gemergde tabelcellen** (``gridSpan`` of ``vMerge``): python-docx's
   ``row.cells`` retourneert dezelfde ``<w:tc>`` meerdere keren via
   verschillende wrappers. Vroeger leidde dat tot dubbele
   block-registratie en daardoor tot mangled placeholders zoals
   ``NL_PHONE_NUMBER_1MBER_1`` na een tweede vervanging op de reeds
   gemoduleerde tekst. We dedupliceren nu op ``id(cell._tc)``.

2. **Hyperlinks** met de PII zelf als display text (Word zet
   ``mailto:foo@bar.nl`` automatisch in een ``<w:hyperlink>`` met de
   email als zichtbare tekst). Voorheen bleef het origineel naast de
   placeholder staan omdat ``para.runs`` alleen top-level runs
   teruggeeft. We strippen nu ook hyperlink-children weg.

Caveats:
    * Velden en comments worden niet aangeraakt.
    * Bookmarks en kruisverwijzingen binnen een vervangen paragraaf
      gaan verloren. Voor de pilot acceptabel.
"""

from __future__ import annotations

import io
from collections.abc import Iterable
from typing import TYPE_CHECKING

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from ._common import apply_replacements_to_text, group_replacements_per_block

# Word-namespace voor lxml-queries op `<w:tc>` etc.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TC_TAG = f"{{{_W_NS}}}tc"

if TYPE_CHECKING:  # pragma: no cover
    from . import AcceptedReplacement, Block, ExtractResult

_BLOCK_SEPARATOR = "\n\n"


def _container_element(container: DocumentType | _Cell):  # type: ignore[no-untyped-def]
    """Geef het XML-element terug waarvan we direct de children willen.

    * Document → ``<w:body>`` (hier leven paragrafen/tabellen).
    * _Cell → ``<w:tc>`` element (cell container).
    * Anders → probeer ``_element`` (protected API, maar stabiel
      genoeg voor python-docx — de ``.element`` property heeft niet
      elke subclass).
    """

    if isinstance(container, DocumentType):
        return container.element.body
    if isinstance(container, _Cell):
        return container._tc
    return container._element


def _iter_container(
    container: DocumentType | _Cell, prefix: str, kind: str = "paragraph"
) -> Iterable[tuple[str, str, Paragraph]]:
    """Loop sequentieel over paragrafen en tabellen in een container.

    Werkt voor zowel een Document als een _Cell. We gebruiken het XML
    rechtstreeks om paragrafen en tabellen in originele volgorde op te
    leveren; dat geeft stabielere id's dan ``container.paragraphs`` +
    ``container.tables`` apart af te lopen.

    Voor tabellen lopen we direct over de fysieke ``<w:tc>`` XML-
    elementen (in plaats van ``row.cells``) zodat een cel met
    ``gridSpan`` of ``vMerge`` slechts één keer wordt bezocht. Voorheen
    leverde ``row.cells`` "virtuele" cell-wrappers op die naar
    hetzelfde ``<w:tc>`` wezen voor gemergde kolommen — daardoor werd
    dezelfde paragraaf-tekst meerdere keren in flat_text geplaatst en
    werden replacements dubbel uitgevoerd (mangled output zoals
    ``NL_PHONE_NUMBER_1MBER_1``).
    """

    element = _container_element(container)
    para_idx = 0
    table_idx = 0
    for child in element.iterchildren():
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            yield (f"{prefix}.p{para_idx}", kind, Paragraph(child, container))
            para_idx += 1
        elif tag == "tbl":
            table = Table(child, container)
            for row_idx, row in enumerate(table.rows):
                tr_element = row._tr
                for cell_idx, tc in enumerate(tr_element.findall(_TC_TAG)):
                    cell = _Cell(tc, table)
                    yield from _iter_container(
                        cell,
                        f"{prefix}.t{table_idx}.r{row_idx}.c{cell_idx}",
                        kind="table-cell",
                    )
            table_idx += 1


def _iter_all_paragraphs(doc: DocumentType) -> Iterable[tuple[str, str, Paragraph]]:
    """Yield (block_id, kind, paragraph) voor alle tekstvelden in doc.

    Volgorde: per sectie header → body (in document-volgorde, inclusief
    tabellen) → per sectie footer. De id-structuur is deterministisch
    voor hetzelfde document.
    """

    for section_idx, section in enumerate(doc.sections):
        for para_idx, para in enumerate(section.header.paragraphs):
            yield (f"h{section_idx}.p{para_idx}", "header", para)

    yield from _iter_container(doc, "b")

    for section_idx, section in enumerate(doc.sections):
        for para_idx, para in enumerate(section.footer.paragraphs):
            yield (f"f{section_idx}.p{para_idx}", "footer", para)


def extract_docx(file_bytes: bytes) -> ExtractResult:
    """Bouw flat_text en block_map op voor een DOCX-bestand."""

    from . import Block, ExtractResult

    doc = Document(io.BytesIO(file_bytes))
    texts: list[str] = []
    blocks: list[Block] = []
    cursor = 0

    for block_id, kind, para in _iter_all_paragraphs(doc):
        text = para.text
        start = cursor
        end = start + len(text)
        blocks.append(Block(id=block_id, kind=kind, start=start, end=end))  # type: ignore[arg-type]
        texts.append(text)
        cursor = end + len(_BLOCK_SEPARATOR)

    flat_text = _BLOCK_SEPARATOR.join(texts)
    return ExtractResult(flat_text=flat_text, blocks=blocks)


_TEXT_PRODUCING_TAGS = frozenset({"r", "hyperlink", "smartTag", "fldSimple"})


def _replace_paragraph_text(para: Paragraph, new_text: str) -> None:
    """Vervang de tekst van een paragraaf met een enkele plain run.

    Strategie: bewaar ``<w:pPr>`` (paragraaf-stijl: alignment, spacing,
    lettertype-default) en verwijder alle tekst-producerende children
    (runs, hyperlinks, smartTags, fldSimple). Voeg vervolgens één
    nieuwe run toe met ``new_text``.

    Verlies: inline run-formatting (bold/cursief mid-zin) en hyperlinks
    op de oorspronkelijke tekst. Behouden: paragraaf-stijl, alignment,
    table-cell-context.

    Waarom destructief? Voorheen pakten we alleen ``runs[0]`` als anker
    en verwijderden we ``runs[1:]``. Dat liet ``<w:hyperlink>``-children
    intact (die staan náást ``<w:r>`` in de XML, niet erin). Resultaat:
    bij een email die door Word in een hyperlink is gezet, bleef het
    origineel naast de placeholder staan (zie
    ``tests/test_docx_support.py::test_hyperlink_email_replaced``).
    """

    p_element = para._element

    # Verzamel children om te slopen: alles wat tekst produceert.
    # `<w:pPr>` (paragraph properties) en bookmark/proof-error markers
    # blijven staan zodat alinea-stijl en cursor-bookmarks behouden zijn.
    to_remove = []
    for child in p_element:
        tag = child.tag.split("}", 1)[-1]
        if tag in _TEXT_PRODUCING_TAGS:
            to_remove.append(child)

    for child in to_remove:
        p_element.remove(child)

    para.add_run(new_text)


def apply_docx(
    file_bytes: bytes,
    replacements: list[AcceptedReplacement],
    blocks: list[Block],
    footer_note: str | None = None,
) -> bytes:
    """Pas vervangingen toe en geef nieuwe DOCX-bytes terug."""

    doc = Document(io.BytesIO(file_bytes))
    grouped = group_replacements_per_block(blocks, replacements)

    # Bouw een index van block_id → paragraph-ref in de huidige parse.
    paragraphs: dict[str, Paragraph] = {}
    for block_id, _kind, para in _iter_all_paragraphs(doc):
        paragraphs[block_id] = para

    for block_id, block_replacements in grouped.items():
        if not block_replacements:
            continue
        target = paragraphs.get(block_id)
        if target is None:
            # Block uit de extract-stap bestaat niet meer in dit document —
            # defensief overslaan.
            continue
        new_text = apply_replacements_to_text(target.text, block_replacements)
        if new_text != target.text:
            _replace_paragraph_text(target, new_text)

    if footer_note:
        # Een lichte horizontale lijn + cursief italic paragraaf onderaan
        # de body. We schrijven 'm in het body-element zodat hij in elke
        # standaard DOCX-viewer zichtbaar is, ook als er secties of
        # kolommen zijn.
        doc.add_paragraph()  # spacer
        para = doc.add_paragraph()
        run = para.add_run(footer_note)
        run.italic = True
        run.font.size = None  # erf van paragraaf-stijl

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


__all__ = ["apply_docx", "extract_docx"]
