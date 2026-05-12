"""Regressie-tests voor de DOCX-pijplijn.

Dekt de twee bug-klassen die in de pilot opdoken (zie de docstring van
``docx_support`` voor de oorspronkelijke symptomen):

1. **Merged tabel-cellen** (``gridSpan``): dezelfde ``<w:tc>`` wordt door
   ``row.cells`` meerdere keren teruggegeven, wat voorheen tot dubbele
   block-registratie en mangled placeholders leidde
   (``NL_PHONE_NUMBER_1MBER_1`` na een tweede toepassing op de reeds
   gemoduleerde tekst).
2. **Email in hyperlink**: ``para.text`` retourneert hyperlink-display-
   tekst, maar ``para.runs`` niet de hyperlink-interne runs. Bij replace
   bleef het origineel staan náást de placeholder.
"""

from __future__ import annotations

import io

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from lxml import etree

from pii_engine.documents import AcceptedReplacement
from pii_engine.documents.docx_support import _TC_TAG, _W_NS, apply_docx, extract_docx

_W_T_TAG = f"{{{_W_NS}}}t"


def _docx_bytes(doc: DocumentType) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_grid_span(cell, span: int) -> None:
    """Markeer ``cell`` als horizontaal-merged over ``span`` kolommen."""
    tc_pr = cell._tc.get_or_add_tcPr()
    grid_span = etree.SubElement(tc_pr, qn("w:gridSpan"))
    grid_span.set(qn("w:val"), str(span))


def _add_hyperlink_run(paragraph, text: str, rid: str = "rId99") -> None:
    """Voeg een ``<w:hyperlink>`` met run ``text`` toe aan ``paragraph``.

    De ``rId`` hoeft niet te resolven voor onze test — het gaat alleen om de
    XML-structuur, niet om navigeerbaarheid.
    """
    p_element = paragraph._element
    hyperlink = etree.SubElement(p_element, qn("w:hyperlink"))
    hyperlink.set(qn("r:id"), rid)
    run = etree.SubElement(hyperlink, qn("w:r"))
    t = etree.SubElement(run, qn("w:t"))
    t.text = text


# --------------------------------------------------------------------------- #
# Merged-cell regressie
# --------------------------------------------------------------------------- #


class TestMergedCells:
    """gridSpan-merged cellen mogen niet leiden tot dubbele verwerking.

    Belangrijke nuance over python-docx: bij gridSpan=2 op cell B wordt
    ``row.cells[1]._tc`` en ``row.cells[2]._tc`` *hetzelfde* ``<w:tc>``-
    element — dat is precies de dubbele-visit die de productie-bug
    veroorzaakte. We laten alle physical ``<w:tc>``-elementen staan en
    vertrouwen op de dedupe-logica in ``_iter_container``.
    """

    def _build_doc_with_merged_cell(self, content: str) -> bytes:
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        row = table.rows[0]
        row.cells[0].text = "Label"
        cell_b = row.cells[1]
        cell_b.text = content
        _set_grid_span(cell_b, 2)
        return _docx_bytes(doc)

    def test_grid_span_produces_duplicate_tc_wrappers(self):
        """Sanity-check op het bug-mechanisme zelf: zonder onze dedupe-fix
        zou ``row.cells`` dezelfde ``<w:tc>`` meerdere keren leveren.
        """
        data = self._build_doc_with_merged_cell("06-12345678")
        doc = Document(io.BytesIO(data))
        row = doc.tables[0].rows[0]
        tc_ids = [id(c._tc) for c in row.cells]
        # Cell B en C delen dezelfde _tc — dat is de duplicate die we
        # in _iter_container moeten skippen.
        assert tc_ids.count(tc_ids[1]) > 1, f"verwachtte duplicate _tc bij gridSpan, kreeg {tc_ids}"

    def test_merged_cell_visited_only_once(self):
        """Dedupe-fix: gemergde paragraaf staat exact één keer in flat_text."""
        data = self._build_doc_with_merged_cell("06-12345678")
        result = extract_docx(data)
        assert result.flat_text.count("06-12345678") == 1, (
            f"flat_text bevat het nummer meer dan één keer — dedupe werkt niet: "
            f"{result.flat_text!r}"
        )
        block_ids = [b.id for b in result.blocks]
        assert len(block_ids) == len(set(block_ids)), "duplicate block_ids — dedupe werkt niet"

    def test_merged_cell_replacement_not_mangled(self):
        """De originele productie-bug: 06-12345678 → NL_PHONE_NUMBER_1MBER_1.

        Na de fix moet de cel exact ``NL_PHONE_NUMBER_1`` bevatten.
        """
        data = self._build_doc_with_merged_cell("06-12345678")
        result = extract_docx(data)

        idx = result.flat_text.index("06-12345678")
        replacements = [
            AcceptedReplacement(
                start=idx,
                end=idx + len("06-12345678"),
                replacement="NL_PHONE_NUMBER_1",
                original="06-12345678",
            )
        ]
        out_bytes = apply_docx(data, replacements, result.blocks)

        out_doc = Document(io.BytesIO(out_bytes))
        cell_texts = [
            cell.text for table in out_doc.tables for row in table.rows for cell in row.cells
        ]

        # De content-cel bevat de placeholder exact één keer en geen
        # mangled varianten zoals "NL_PHONE_NUMBER_1MBER_1".
        target_cells = [t for t in cell_texts if "NL_PHONE_NUMBER" in t]
        assert target_cells, f"placeholder niet in output: {cell_texts!r}"
        for txt in target_cells:
            # Exact-match check — geen suffix-fragment, geen origineel,
            # geen duplicaat.
            assert txt == "NL_PHONE_NUMBER_1", f"mangled placeholder in cel: {txt!r}"


# --------------------------------------------------------------------------- #
# Hyperlink regressie
# --------------------------------------------------------------------------- #


class TestHyperlinkEmail:
    """Een email die in een <w:hyperlink> staat moet volledig vervangen worden."""

    def _build_doc_with_email_hyperlink(self, email: str) -> bytes:
        doc = Document()
        para = doc.add_paragraph()
        _add_hyperlink_run(para, email)
        return _docx_bytes(doc)

    def test_hyperlink_email_visible_in_flat_text(self):
        data = self._build_doc_with_email_hyperlink("josverstappen@hotmail.nl")
        result = extract_docx(data)
        assert "josverstappen@hotmail.nl" in result.flat_text
        # Niet dubbel — moet exact één keer voorkomen.
        assert result.flat_text.count("josverstappen@hotmail.nl") == 1

    def test_hyperlink_email_replaced_cleanly(self):
        email = "josverstappen@hotmail.nl"
        data = self._build_doc_with_email_hyperlink(email)
        result = extract_docx(data)

        idx = result.flat_text.index(email)
        replacements = [
            AcceptedReplacement(
                start=idx,
                end=idx + len(email),
                replacement="EMAIL_ADDRESS_1",
                original=email,
            )
        ]
        out_bytes = apply_docx(data, replacements, result.blocks)

        out_doc = Document(io.BytesIO(out_bytes))
        body_text = "\n".join(p.text for p in out_doc.paragraphs)

        # Origineel mag nergens meer staan.
        assert email not in body_text, f"origineel email staat nog in output: {body_text!r}"
        # Placeholder moet er één keer in staan.
        assert body_text.count("EMAIL_ADDRESS_1") == 1, f"placeholder count != 1: {body_text!r}"

        # Extra: er mag geen <w:hyperlink> meer in de paragraaf zitten.
        para = out_doc.paragraphs[0]
        hyperlinks = para._element.findall(qn("w:hyperlink"))
        assert hyperlinks == [], "hyperlink-element bleef bestaan na replace"


# --------------------------------------------------------------------------- #
# Combinatie van beide bugs in één paragraaf
# --------------------------------------------------------------------------- #


class TestMergedCellPlusHyperlink:
    """De exacte combinatie uit het HAN-formulier dat de bug aan het licht
    bracht: 3-koloms tabel met een merged content-cel (gridSpan=2) die een
    email-hyperlink bevat."""

    def test_merged_cell_with_hyperlink_email_no_leak(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "Label"
        table.rows[0].cells[1].text = "Nieuwe gegevens"
        table.rows[0].cells[2].text = "Oude gegevens"

        table.rows[1].cells[0].text = "E-mailadres"
        email_cell = table.rows[1].cells[1]
        email_para = email_cell.paragraphs[0]
        _add_hyperlink_run(email_para, "josverstappen@hotmail.nl")
        _set_grid_span(email_cell, 2)

        data = _docx_bytes(doc)
        result = extract_docx(data)

        # Email moet exact één keer in flat_text staan — niet drie (1× per
        # bezochte cell × hyperlink+plain-run varianten).
        assert result.flat_text.count("josverstappen@hotmail.nl") == 1, (
            f"email staat meerdere keren in flat_text: {result.flat_text!r}"
        )

        idx = result.flat_text.index("josverstappen@hotmail.nl")
        replacements = [
            AcceptedReplacement(
                start=idx,
                end=idx + len("josverstappen@hotmail.nl"),
                replacement="EMAIL_ADDRESS_1",
                original="josverstappen@hotmail.nl",
            )
        ]
        out_bytes = apply_docx(data, replacements, result.blocks)

        out_doc = Document(io.BytesIO(out_bytes))

        # Verzamel tekst per fysieke <w:tc> (niet via row.cells, die levert
        # virtuele duplicaten bij gridSpan op).
        physical_cell_texts: list[str] = []
        for tbl in out_doc.tables:
            for row in tbl.rows:
                for tc in row._tr.findall(_TC_TAG):
                    text = "".join(t.text or "" for t in tc.iter(_W_T_TAG))
                    physical_cell_texts.append(text)

        joined = " | ".join(physical_cell_texts)

        # Origineel email-adres mag nergens meer in de output staan.
        assert "josverstappen@hotmail.nl" not in joined, (
            f"origineel email lekt naar output: {joined!r}"
        )
        # Geen mangled of dubbele placeholders.
        assert "EMAIL_ADDRESS_1EMAIL_ADDRESS_1" not in joined
        # Email-placeholder zit precies in één fysieke cell.
        assert sum(1 for t in physical_cell_texts if "EMAIL_ADDRESS_1" in t) == 1, (
            f"placeholder zit niet in exact 1 fysieke cel: {physical_cell_texts!r}"
        )
