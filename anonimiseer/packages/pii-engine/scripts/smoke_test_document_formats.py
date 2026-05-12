"""End-to-end smoke-test voor de Anonimiseer document-pijplijn.

Genereert drie synthetische test-documenten met identieke PII-inhoud in
docx / xlsx / pdf, pompt ze door de lokale engine
(``http://127.0.0.1:8765``), en checkt dat geen origineel meer in de
output staat én dat er geen mangled placeholders zijn.

Bedoeld als **dev-tool**, niet als pytest-suite — vereist een lopende
engine. De unit-tests in ``tests/test_docx_support.py`` dekken dezelfde
bugs geautomatiseerd zonder engine.

Gebruik::

    # Engine alvast draaien:
    pii-engine &

    # Smoke-test:
    python scripts/smoke_test_document_formats.py [--keep-outputs]

Met ``--keep-outputs`` blijven de in/uit-bestanden in de huidige map
staan zodat je ze handmatig kunt openen in Word/Excel/Preview.

Achtergrond: deze suite is ontstaan na de v0.1.0-bug waarbij DOCX-
formulieren met gemergde tabel-cellen en hyperlinks tot mangled
placeholders (``NL_PHONE_NUMBER_1MBER_1``) en zichtbaar gebleven
originelen leidden. Zie ``tests/test_docx_support.py`` voor de
geautomatiseerde regressie-tests.
"""

from __future__ import annotations

import argparse
import io
import json
import sys
from pathlib import Path

import httpx
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from openpyxl import Workbook
from openpyxl.worksheet.hyperlink import Hyperlink
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

API = "http://127.0.0.1:8765"

# Eén centrale dataset zodat alle 3 de formaten dezelfde PII bevatten en
# we de output naast elkaar kunnen leggen.
PII = {
    "naam": "Jos Verstappen",
    "telefoon": "06-12345678",
    "telefoon_aanspreekpunt": "+31 6 87654321",
    "email": "josverstappen@hotmail.nl",
    "email_aanspreekpunt": "han.goedkeurder@example.com",
    "bsn": "111222333",  # Elfproef-geldig test-BSN
    "iban": "NL91ABNA0417164300",
    "adres": "Ruitenberglaan 26, 6826 CC Arnhem",
    "geboortedatum": "12-03-1985",
    "studentnummer": "s1234567",
}

# Patronen die in de geanonimiseerde output GEEN substring meer mogen
# zijn. Voor de placeholder-mangling testen we op specifieke buggy
# patronen die we eerder zagen (en op dubbele placeholders).
ORIGINALS_THAT_MUST_DISAPPEAR = [
    PII["email"],
    PII["email_aanspreekpunt"],
    PII["telefoon"],
    PII["telefoon_aanspreekpunt"],
    PII["bsn"],
    PII["iban"],
]

MANGLED_PATTERNS = [
    "NL_PHONE_NUMBER_1MBER_1",
    "NL_PHONE_NUMBER_2UMBER_2",
    "EMAIL_ADDRESS_1EMAIL_ADDRESS_1",
    "EMAIL_ADDRESS_2EMAIL_ADDRESS_2",
    "PERSON_1ERSON_1",
    "LOCATION_1ON_1",
]


# --------------------------------------------------------------------------- #
# Fixture builders
# --------------------------------------------------------------------------- #


def _add_hyperlink_run(paragraph, text: str, rid: str = "rIdMail") -> None:
    """Plaats een <w:hyperlink>-element met run(text) in een DOCX-paragraaf."""
    hyperlink = etree.SubElement(paragraph._element, qn("w:hyperlink"))
    hyperlink.set(qn("r:id"), rid)
    run = etree.SubElement(hyperlink, qn("w:r"))
    t = etree.SubElement(run, qn("w:t"))
    t.text = text


def _set_grid_span(cell, span: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    grid_span = etree.SubElement(tc_pr, qn("w:gridSpan"))
    grid_span.set(qn("w:val"), str(span))


def build_docx() -> bytes:
    """DOCX die de bekende HAN-formulier-edge cases nabootst.

    Bevat:
      - Plain alinea met persoonsnaam, BSN, IBAN.
      - 3-koloms tabel met merged content-cellen (gridSpan=2).
      - Email-cel met <w:hyperlink>.
    """
    doc = Document()
    doc.add_heading("Inzet HANFlex-er (CIM formulier)", level=1)
    doc.add_paragraph(
        f"Kandidaat: {PII['naam']}, geboren op {PII['geboortedatum']}, "
        f"studentnummer {PII['studentnummer']}, BSN {PII['bsn']}. "
        f"Adres: {PII['adres']}. IBAN: {PII['iban']}."
    )

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"

    table.rows[0].cells[0].text = "Omschrijving"
    table.rows[0].cells[1].text = "Nieuwe gegevens"
    table.rows[0].cells[2].text = "Oude gegevens"

    rows = [
        ("Naam kandidaat", PII["naam"]),
        ("Telefoonnummer", PII["telefoon"]),
        ("E-mailadres", PII["email"]),
        ("Geboortedatum", PII["geboortedatum"]),
        ("Aanspreekpunt", "Anna Jansen"),
    ]
    for r, (label, value) in enumerate(rows, start=1):
        table.rows[r].cells[0].text = label
        content_cell = table.rows[r].cells[1]
        if label == "E-mailadres":
            _add_hyperlink_run(content_cell.paragraphs[0], value)
        else:
            content_cell.text = value
        # Merge de "Nieuwe gegevens"-cel met "Oude gegevens" → gridSpan=2.
        _set_grid_span(content_cell, 2)

    doc.add_paragraph(
        f"Aanspreekpunt telefoonnummer: {PII['telefoon_aanspreekpunt']}, "
        f"email: {PII['email_aanspreekpunt']}."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_xlsx() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Inzet"
    ws["A1"] = "Veld"
    ws["B1"] = "Nieuwe gegevens"
    ws["C1"] = "Oude gegevens"

    rows = [
        ("Naam kandidaat", PII["naam"]),
        ("Telefoonnummer", PII["telefoon"]),
        ("E-mailadres", PII["email"]),
        ("Geboortedatum", PII["geboortedatum"]),
        ("BSN", PII["bsn"]),
        ("IBAN", PII["iban"]),
        ("Adres", PII["adres"]),
        ("Studentnummer", PII["studentnummer"]),
        ("Aanspreekpunt naam", "Anna Jansen"),
        ("Aanspreekpunt email", PII["email_aanspreekpunt"]),
    ]
    for idx, (label, value) in enumerate(rows, start=2):
        ws.cell(row=idx, column=1, value=label)
        cell_b = ws.cell(row=idx, column=2, value=value)
        if "email" in label.lower():
            cell_b.hyperlink = Hyperlink(ref=cell_b.coordinate, target=f"mailto:{value}")

    # Een paar gemergde cellen voor realisme
    ws.merge_cells("B12:C12")
    ws["A12"] = "Toelichting"
    ws["B12"] = f"Contact via {PII['email_aanspreekpunt']} of {PII['telefoon_aanspreekpunt']}."

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_pdf() -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    _width, height = A4
    c.setFont("Helvetica-Bold", 14)
    c.drawString(20 * mm, height - 25 * mm, "Inzet HANFlex-er (CIM formulier)")
    c.setFont("Helvetica", 11)
    y = height - 40 * mm
    lines = [
        f"Naam kandidaat: {PII['naam']}",
        f"Geboortedatum: {PII['geboortedatum']}",
        f"Studentnummer: {PII['studentnummer']}",
        f"BSN: {PII['bsn']}",
        f"IBAN: {PII['iban']}",
        f"Adres: {PII['adres']}",
        f"Telefoonnummer: {PII['telefoon']}",
        f"E-mailadres: {PII['email']}",
        "",
        "Aanspreekpunt:",
        "  Naam: Anna Jansen",
        f"  Telefoon: {PII['telefoon_aanspreekpunt']}",
        f"  Email: {PII['email_aanspreekpunt']}",
    ]
    for line in lines:
        c.drawString(20 * mm, y, line)
        y -= 7 * mm
    c.save()
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Engine-driver
# --------------------------------------------------------------------------- #


def anonymize_via_api(filename: str, data: bytes) -> tuple[bytes, dict]:
    """Loop een bestand door de engine: extract → analyze → apply.

    Retourneert de geanonimiseerde bytes plus diagnostiek (aantal blocks,
    aantal hits, sample replacements).
    """
    with httpx.Client(base_url=API, timeout=120.0) as client:
        r = client.post(
            "/document/extract",
            files={"file": (filename, data, "application/octet-stream")},
        )
        r.raise_for_status()
        extract = r.json()

        r = client.post(
            "/analyze",
            json={"text": extract["flat_text"], "language": "nl"},
        )
        r.raise_for_status()
        analysis = r.json()
        hits = analysis.get("items") or analysis.get("results") or []

        # Maak pseudoniem-replacements (één ID per (entity_type, original_text)).
        counters: dict[str, int] = {}
        canonical: dict[tuple[str, str], str] = {}
        replacements = []
        for hit in hits:
            etype = hit["entity_type"]
            start = hit["start"]
            end = hit["end"]
            original = extract["flat_text"][start:end]
            key = (etype, original.lower())
            if key not in canonical:
                counters[etype] = counters.get(etype, 0) + 1
                canonical[key] = f"{etype}_{counters[etype]}"
            replacement = canonical[key]
            replacements.append(
                {
                    "start": start,
                    "end": end,
                    "replacement": replacement,
                    "original": original,
                }
            )

        payload = {
            "replacements": replacements,
            "blocks": extract["blocks"],
            "footer_note": None,
        }
        r = client.post(
            "/document/apply",
            files={"file": (filename, data, "application/octet-stream")},
            data={"payload": json.dumps(payload)},
        )
        r.raise_for_status()
        return r.content, {
            "n_blocks": len(extract["blocks"]),
            "n_hits": len(hits),
            "n_replacements": len(replacements),
            "sample_replacements": replacements[:6],
        }


# --------------------------------------------------------------------------- #
# Output-extractie voor verificatie
# --------------------------------------------------------------------------- #


def read_docx_text(data: bytes) -> str:
    from pii_engine.documents.docx_support import extract_docx  # type: ignore

    return extract_docx(data).flat_text


def read_xlsx_text(data: bytes) -> str:
    from pii_engine.documents.xlsx_support import extract_xlsx  # type: ignore

    return extract_xlsx(data).flat_text


def read_pdf_text(data: bytes) -> str:
    from pii_engine.documents.pdf_support import extract_pdf  # type: ignore

    return extract_pdf(data).flat_text


# --------------------------------------------------------------------------- #
# Test-runner
# --------------------------------------------------------------------------- #


def run_case(label: str, filename: str, data: bytes, read_text) -> dict:
    print(f"\n── {label} ({filename}, {len(data):,} bytes) ──")
    try:
        out_bytes, diag = anonymize_via_api(filename, data)
    except Exception as e:
        print(f"  ❌ ENGINE-FOUT: {e}")
        return {"label": label, "ok": False, "error": str(e)}

    print(f"  blocks={diag['n_blocks']}, hits={diag['n_hits']}")
    out_text = read_text(out_bytes)

    failures = []
    for needle in ORIGINALS_THAT_MUST_DISAPPEAR:
        if needle in out_text:
            failures.append(f"LEAK: origineel {needle!r} staat nog in output")
    for needle in MANGLED_PATTERNS:
        if needle in out_text:
            failures.append(f"MANGLED placeholder: {needle!r}")

    if failures:
        print("  ❌ FAIL:")
        for f in failures:
            print(f"     · {f}")
    else:
        print("  ✓ origineel weg, geen mangled placeholders")
    return {
        "label": label,
        "filename": filename,
        "ok": not failures,
        "failures": failures,
        "n_blocks": diag["n_blocks"],
        "n_hits": diag["n_hits"],
        "out_bytes": out_bytes,
        "out_text_preview": out_text[:400],
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--keep-outputs",
        action="store_true",
        help="Bewaar gegenereerde input- en outputbestanden in de huidige map.",
    )
    args = parser.parse_args()

    docx_bytes = build_docx()
    xlsx_bytes = build_xlsx()
    pdf_bytes = build_pdf()

    results = []
    results.append(run_case("DOCX", "fixture.docx", docx_bytes, read_docx_text))
    results.append(run_case("XLSX", "fixture.xlsx", xlsx_bytes, read_xlsx_text))
    results.append(run_case("PDF", "fixture.pdf", pdf_bytes, read_pdf_text))

    if args.keep_outputs:
        outdir = Path.cwd()
        for r, in_bytes, ext in (
            (results[0], docx_bytes, ".docx"),
            (results[1], xlsx_bytes, ".xlsx"),
            (results[2], pdf_bytes, ".pdf"),
        ):
            (outdir / f"input{ext}").write_bytes(in_bytes)
            (outdir / f"output{ext}").write_bytes(r["out_bytes"])
        print(f"\nbestanden in {outdir}")

    print("\n── SAMENVATTING ──")
    ok = sum(1 for r in results if r["ok"])
    print(f"{ok}/{len(results)} formaten OK")
    for r in results:
        flag = "✓" if r["ok"] else "✗"
        print(f"  {flag} {r['label']}")
    return 0 if ok == len(results) else 1


if __name__ == "__main__":
    sys.exit(main())
